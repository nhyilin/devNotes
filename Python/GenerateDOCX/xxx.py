from docx import Document
from docx.shared import Cm
from docx.oxml.shared import OxmlElement, qn


def set_cell_margins(cell, **kwargs):
    '''设置某单元格间距

    长度单位为Twips，1Twips = 1/20pt，1Twips = 1/567cm

    >>> set_cell_margins(table.cell(1, 0), top=0, start=0, bottom=0, end=0, left=0, right=0)

    :param cell: 某单元格
    :param top: 上边距
    :param start: 左边距
    :param bottom: 下边距
    :param end: 右边距
    :param left: 左边距（WPS）
    :param right: 右边距（WPS）
    '''
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m in ['top', 'start', 'bottom', 'end', 'left', 'right']:
        if m in kwargs:
            node = OxmlElement('w:{}'.format(m))
            node.set(qn('w:w'), str(kwargs.get(m)))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
    tcPr.append(tcMar)


document = Document()

table = document.add_table(2, 2, style='Light List')
heading_cells = table.rows[0].cells
cell1, cell2 = heading_cells[0], heading_cells[-1]
cell1.merge(cell2)  # 合并单元格
cell1.text = '示例图像'
table.rows[1].height = Cm(5.0)  # 第二行行高为5cm

detail_cell = table.cell(1, 0)  # 描述图片的单元格
picture_cell = table.cell(1, 1)  # 放置图片的单元格
detail_cell.text = '这是一只边牧'
paragraph = picture_cell.paragraphs[0]
run = paragraph.add_run()
picture = run.add_picture('test.jpg', height=Cm(5.0))  # 插入图片
picture_cell.width = Cm(picture.width.cm)  # 单元格宽度设为图片的宽度

min_width = min([cell.width.cm for cell in table.column_cells(1)])  # 第二列单元格的最小宽度

for cell in table.column_cells(1):  # 第二列单元格
    cell.width = Cm(min_width)  # 第二列单元格宽度设为最小宽度

set_cell_margins(picture_cell, top=0, start=0, bottom=0, end=0, left=0, right=0)  # 设置单元格间距

document.save('test.docx')
