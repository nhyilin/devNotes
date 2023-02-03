# -*- coding : utf-8-*-
import time as t
import threading
from docx.shared import Inches, Pt
from docxtpl import DocxTemplate
import docx
import pandas as pd
import os
import csv
import numpy as np
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import cn2an

CURRENT_DIR = os.getcwd() + '/'  # 获取当前的路径
TEMPFILE = "情报要报.docx"
CSV_PATH = os.getcwd() + "/data/"
NEW_FILE_PATH = os.getcwd() + "/new/"


def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def mkdir(path):
    path = path.strip()
    isExists = os.path.exists(path)
    if not isExists:
        os.makedirs(path.encode('utf-8'))
        return path
    else:
        return path


def list_dir(file_dir):
    list_csv = []
    dir_list = os.listdir(file_dir)
    for cur_file in dir_list:
        path = os.path.join(file_dir, cur_file)
        if os.path.isfile(path):
            dir_files = os.path.join(file_dir, cur_file)
        if os.path.splitext(path)[1] == '.csv':
            csv_file = os.path.join(file_dir, cur_file)
            list_csv.append(csv_file)
        if os.path.isdir(path):
            list_dir(path)
    return list_csv


def list_word_dir(file_dir):
    list_csv = []
    dir_list = os.listdir(file_dir)
    for cur_file in dir_list:
        path = os.path.join(file_dir, cur_file)
        if os.path.isfile(path):
            dir_files = os.path.join(file_dir, cur_file)
        if os.path.splitext(path)[1] == '.docx':
            csv_file = os.path.join(file_dir, cur_file)
            list_csv.append(csv_file)
        if os.path.isdir(path):
            list_dir(path)
    return list_csv


def makeDic(key, value):
    dic = {key: value}
    return dic


def EmptyDir(filepath):
    # 清空目标文件夹
    del_list = os.listdir(filepath)
    for f in del_list:
        file_path = os.path.join(filepath, f)
        if os.path.isfile(file_path):
            os.remove(file_path)


def numToChineseWord(str):
    output = cn2an.an2cn(str) + "、"
    return output


def ReFormatColumn(data, number, textContent, row, key):
    """
    data：字典中的值，将很多条目合并到一起的目标，入值最好为空字符串
    number：csv文件中内容条数索引
    content：csv文件中的内容，如biref和Content
    row:csv文件行数，以供遍历
    key:csv文件的键
    """
    for i in range(row):
        data += numToChineseWord(str(number[i]))
        # data += '. '
        # data += ' '
        data += str(textContent[i]).replace("xAyBzC01", "\n    ")
        data += '\n    '
    return makeDic(key, data)


def FormatCSV(csvfile):
    global DIC
    DIC = dict()

    global time
    data = pd.read_csv(csvfile)
    # encoding = 'gbk'
    time = data["Time"]
    rerange = data["Issue"]
    number = data["Item"]
    Type = data["Type"]
    Importance = data["Importance"]
    Brief = data["Brief"]
    Content = data["Content"]
    Publisher = data["Publisher"]
    SourceUri = data["SourceUri"]
    # TABLE_PATH 是后文中贴的表单的路径列表
    TABLE_PATH = []
    for value in SourceUri:
        TABLE_PATH.append(str(value))

    num = data.shape[0]  # 获取数据行数

    global Brief_dic
    Brief_data = str()
    Brief_dic = ReFormatColumn(Brief_data, number, Brief, num, "Brief")
    DIC.update(Brief_dic)

    global Content_dic
    Content_data = str()
    Content_dic = ReFormatColumn(Content_data, number, Content, num, "Content")
    DIC.update(Content_dic)

    Time_dic = {"Time": time[0]}
    DIC.update(Time_dic)

    Rerange_dic = {"期号": rerange[0]}
    DIC.update(Rerange_dic)

    End_dic = {"此页无正文": "（此页无正文）"}
    DIC.update(End_dic)
    # print("Final dictionary: ", '\n', DIC)
    return TABLE_PATH


def delete_paragraph(paragraph):
    # 清除word中多余空行
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def DeleteBlankPages(xxx):
    # 清除word中空白页面
    myDoc = Document(xxx)
    for num, paragraphs in enumerate(myDoc.paragraphs):
        if paragraphs.text == "":
            # print('第{}段是空行段'.format(num))
            paragraphs.clear()  # 清除文字，并不删除段落，run也可以,
            delete_paragraph(paragraphs)


"""
# 该函数提供打印word中内容的功能
def PrintWord():
    file = docx.Document(NEW_FILE_PATH)
    print("段落数:" + str(len(file.paragraphs)))  # 段落数为13，每个回车隔离一段

    # 输出段落编号及段落内容
    for i in range(len(file.paragraphs)):
        if len(file.paragraphs[i].text.replace(' ', '')) > 4:
            print("第" + str(i) + "段的内容是：" + file.paragraphs[i].text)

"""


def AddTable(wordFile, csvFile, latestFileName):
    """
    wordFile：要添加表格的word
    csvFile：要添加的表格
    latestFileName：最终生成的word文件的名称
    """
    doc = docx.Document(wordFile)
    index = 0
    for csvIndex in csvFile:
        if not csvIndex:
            continue
        if len(csvIndex) <= 21:
            print("表格路径过短，已跳过，csv中标注路径为：" + csvIndex)
            continue
        if len(csvIndex) > 40:
            print("表格路径过长，已跳过，csv中标注路径为：" + csvIndex)
            continue
        csvIndex = csvIndex[1:]
        with open(csvIndex) as f:
            f_csv = csv.DictReader(f)
            ROWS = 0
            csvKeyInotList = []
            csvValueInotList = []
            for row in f_csv:
                # csv行数，此系数极为重要
                ROWS = ROWS + 1
                for key, value in row.items():
                    csvKeyInotList.append(key)
                    csvValueInotList.append(value)
        del csvKeyInotList[11:]
        csvKeyInotList = np.array(csvKeyInotList)
        csvValueInotList = np.array(csvValueInotList).reshape(ROWS, 11)

        table = doc.add_table(rows=ROWS + 1, cols=11, style='Table Grid')
        table.autofit = True

        hdr_cells = table.rows[0].cells
        hdr_cells[5].width = Inches(1.3)
        hdr_cells[6].width = Inches(2.0)
        hdr_cells[10].width = Inches(1.8)

        for cell in table.columns[0].cells:
            cell.width = Inches(0.2)
        table_header = table.rows[0].cells
        for col_x in range(0, 11):
            table_header[col_x].text = csvKeyInotList[col_x]
            table_header[col_x].paragraphs[0].runs[0].font.size = Pt(12)

        for row in range(1, ROWS + 1):
            for col_y in range(0, 11):
                table.rows[row].cells[col_y].text = csvValueInotList[row - 1][col_y]
                table.rows[row].cells[col_y].paragraphs[0].runs[0].font.size = Pt(12)

        # 首行灰色
        rows = table.rows[0]
        for cell in rows.cells:
            shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
            cell._tc.get_or_add_tcPr().append(shading_elm)

        # print("段落数:" + str(len(doc.paragraphs)))
        # print(doc.paragraphs[0])
        # 此处doc.paragraphs[]为经验值，试出来的

        # move_table_after函数极为重要，提供将新建的table移动到哪里的功能
        move_table_after(table, doc.paragraphs[11 + index])

        table_headers = doc.add_table(1, cols=1)
        cell = table_headers.cell(0, 0)
        cell.line_spacing = 0.5

        table_name = ""
        if csvIndex[29] == "1":
            table_name = "卫星轨道降低幅度较大的统计"
        elif csvIndex[29] == "2":
            table_name = "卫星轨道降低幅度较小的统计"
        elif csvIndex[29] == "3":
            table_name = "卫星轨道升高幅度较小的统计"
        else:
            table_name = "卫星轨道升高幅度较大的统计"

        cell.paragraphs[0].add_run("附表 " + str(index + 1) + table_name).bold = True
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 水平居中

        move_table_after(table_headers, doc.paragraphs[11 + index])

        index = index + 1
        DeleteBlankPages(latestFileName)
        doc.save(latestFileName)


"""
# 该函数提供从csv中解析出table路径的功能，但是由于功能重复，且IO重复占用，暂时取消使用
def getTargetFile(file):
    # 获取要添加的csv文件
    csvList = []
    with open(file) as f:
        csvList = []
        f_csv = csv.DictReader(f)
        for key in f_csv:
            csvList.append(key["SourceUri"])
        return csvList
"""


def ProcessingThread(csv_list, counter):
    for index in csv_list:
        counter = counter + 1
        TABLE_PATH = FormatCSV(index)
        tpl = DocxTemplate(CURRENT_DIR + TEMPFILE)
        tpl.render(DIC)  # 渲染替换
        tpl.save(NEW_FILE_PATH + r"{}.docx".format(time[0]))

        AddTable(NEW_FILE_PATH + r"{}.docx".format(time[0]), TABLE_PATH,
                 NEW_FILE_PATH + r"{}.docx".format(time[0]))

        print("Newly generated files： ", NEW_FILE_PATH + r"{}.docx".format(time[0]))
        DeleteBlankPages(NEW_FILE_PATH + r"{}.docx".format(time[0]))
        print("Total number of documents: " + str(len(csv_list)) + "\t Generated: " + str(counter))


def main():
    tic = t.perf_counter()
    counter = 0
    EmptyDir(NEW_FILE_PATH)
    csv_list = list_dir(CSV_PATH)
    print("csv_fils: ", '\n', csv_list)
    # ProcessingThread(csv_list, counter)
    # for index in csv_list:
    #     counter = counter + 1
    #     TABLE_PATH = FormatCSV(index)
    #     tpl = DocxTemplate(CURRENT_DIR + TEMPFILE)
    #     tpl.render(DIC)  # 渲染替换
    #     tpl.save(NEW_FILE_PATH + r"{}.docx".format(time[0]))
    #
    #     add_table(NEW_FILE_PATH + r"{}.docx".format(time[0]), TABLE_PATH,
    #              NEW_FILE_PATH + r"{}.docx".format(time[0]))
    #
    #     print("Newly generated files： ", NEW_FILE_PATH + r"{}.docx".format(time[0]))
    #     delete_blank_pages(NEW_FILE_PATH + r"{}.docx".format(time[0]))
    #     print("Total number of documents: " + str(len(csv_list)) + "\t Generated: " + str(counter))

    toc = t.perf_counter()
    print(f"time cost: {toc - tic:0.4f} seconds", "\t each file cost:" + str((toc - tic) / counter)[:5])
    predictTime(counter, tic, toc)


def predictTime(counter, tic, toc):
    seconds = ((toc - tic) / counter) * 1000
    m, s = divmod(seconds, 60)
    h, m = divmod(m, 60)
    print("if total number was 1000, it will cost:" + "%02d:%02d:%02d" % (h, m, s))

class MyThread(threading.Thread):
    def __init__(self, thread_name):
        # 注意：一定要显式的调用父类的初始化函数。
        super(MyThread, self).__init__(name=thread_name)

    def run(self):
        ProcessingThread(csv_list, counter)
        print("%s正在运行中......" % self.name)

if __name__ == "__main__":
    main()
