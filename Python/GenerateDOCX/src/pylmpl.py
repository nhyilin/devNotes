# -*- coding : utf-8-*-
import time as t
from docx.shared import Inches, Pt
from docxtpl import DocxTemplate
import docx
import pandas as pd
import os
import csv
import numpy as np
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx import Document
import cn2an
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm
from docx.oxml.shared import OxmlElement, qn
import word2pdf


# CURRENT_DIR = os.getcwd() + '/'  # 获取当前的路径
# TEMPFILE = "新模版.docx"
# CSV_PATH = os.getcwd() + "/rptD_ONE/"
# NEW_FILE_PATH = os.getcwd() + "/output_doc/"


def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def mkdir(path):
    path = path.strip()
    isExists = os.path.exists(path)
    if not isExists:
        os.makedirs(path.encode('utf-8'))
        return path
    else:
        return path


def list_dir(file_dir):
    list_csv = []
    dir_list = os.listdir(file_dir)
    for cur_file in dir_list:
        path = os.path.join(file_dir, cur_file)
        if os.path.isfile(path):
            dir_files = os.path.join(file_dir, cur_file)
        if os.path.splitext(path)[1] == '.csv':
            csv_file = os.path.join(file_dir, cur_file)
            list_csv.append(csv_file)
        if os.path.isdir(path):
            list_dir(path)
    return list_csv


def list_word_dir(file_dir):
    list_csv = []
    dir_list = os.listdir(file_dir)
    for cur_file in dir_list:
        path = os.path.join(file_dir, cur_file)
        if os.path.isfile(path):
            dir_files = os.path.join(file_dir, cur_file)
        if os.path.splitext(path)[1] == '.docx':
            csv_file = os.path.join(file_dir, cur_file)
            list_csv.append(csv_file)
        if os.path.isdir(path):
            list_dir(path)
    return list_csv


def makeDic(key, value):
    dic = {key: value}
    return dic


def EmptyDir(filepath):
    # 清空目标文件夹
    if not os.path.exists(filepath):
        os.makedirs(filepath)
    del_list = os.listdir(filepath)
    for f in del_list:
        file_path = os.path.join(filepath, f)
        if os.path.isfile(file_path):
            os.remove(file_path)


def numToChineseWord(str):
    output = cn2an.an2cn(str) + "、"
    return output


def numToChineseWord2(str):
    output = cn2an.an2cn(str)
    return output


def re_format_column(data, number, textContent, row, key):
    """
    data：字典中的值，将很多条目合并到一起的目标，入值最好为空字符串
    number：csv文件中内容条数索引
    content：csv文件中的内容，如biref和Content
    row:csv文件行数，以供遍历
    key:csv文件的键
    """
    for i in range(row):
        data += numToChineseWord(str(number[i]))
        # data += '. '
        # data += ' '
        data += str(textContent[i]).replace("xAyBzC01", "\n    ")
        data += '\n    '
    tmpD = data.find("xAyBzC02")
    if tmpD > -1:
        numofsecondheader = 1
        for i in range(1, 5, 1):
            data = data.replace("xAyBzC02", "\n    " + "（" + numToChineseWord2(str(numofsecondheader)) + "） ", 1)
            tmpD = data.find("xAyBzC02")
            if tmpD == -1:
                break
            numofsecondheader = numofsecondheader + 1

    return makeDic(key, data)


def FormatCSV(csvfile):
    global DIC
    DIC = dict()

    global time
    data = pd.read_csv(csvfile)
    # encoding = 'gbk'
    time = data["Time"]
    rerange = data["Issue"]
    number = data["Item"]
    Type = data["Type"]
    Importance = data["Importance"]
    Brief = data["Brief"]
    Content = data["Content"]
    Publisher = data["Publisher"]
    SourceUri = data["SourceUri"]
    # TABLE_PATH 是后文中贴的表单的路径列表
    TABLE_PATH = []
    for value in SourceUri:
        TABLE_PATH.append(str(value))

    num = data.shape[0]  # 获取数据行数

    global Brief_dic
    Brief_data = str()
    Brief_dic = re_format_column(Brief_data, number, Brief, num, "Brief")
    DIC.update(Brief_dic)

    global Content_dic
    Content_data = str()
    Content_dic = re_format_column(Content_data, number, Content, num, "Content")
    DIC.update(Content_dic)

    Time_dic = {"Time": time[0]}
    DIC.update(Time_dic)

    Rerange_dic = {"期号": rerange[0]}
    DIC.update(Rerange_dic)

    End_dic = {"此页无正文": "（此页无正文）"}
    DIC.update(End_dic)
    # print("Final dictionary: ", '\n', DIC)
    return TABLE_PATH


def delete_paragraph(paragraph):
    # 清除word中多余空行
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def delete_blank_pages(xxx):
    # 清除word中空白页面
    myDoc = Document(xxx)
    for num, paragraphs in enumerate(myDoc.paragraphs):
        if paragraphs.text == "":
            # print('第{}段是空行段'.format(num))
            paragraphs.clear()  # 清除文字，并不删除段落，run也可以,
            delete_paragraph(paragraphs)


"""
# 该函数提供打印word中内容的功能
def PrintWord():
    file = docx.Document(NEW_FILE_PATH)
    print("段落数:" + str(len(file.paragraphs)))  # 段落数为13，每个回车隔离一段

    # 输出段落编号及段落内容
    for i in range(len(file.paragraphs)):
        if len(file.paragraphs[i].text.replace(' ', '')) > 4:
            print("第" + str(i) + "段的内容是：" + file.paragraphs[i].text)

"""


def multi_table_in_one_chart(csvFileChart):
    """
    检测csv中一个表单项中有多个表格地址，并添加所有列出项
    csvFileChart：一个csv中表格列表，如：
    ['nan', '../Z1_rptD_alert_lsj/20211115-1.csv', '../Z1_rptD_alert_lsj/20211115-4.csv']
    """
    newcsvFile = []
    for urls in csvFileChart:
        urls.replace(" ", "")
        urls = urls.split("xAyBzC03")
        for newUrls in urls:
            newUrls = newUrls.replace(" ", "")
            newcsvFile.append(newUrls)
    return newcsvFile


def set_cell_margins(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m in ['top', 'start', 'bottom', 'end', 'left', 'right']:
        if m in kwargs:
            node = OxmlElement('w:{}'.format(m))
            node.set(qn('w:w'), str(kwargs.get(m)))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
    tcPr.append(tcMar)


def add_table(wordFile, csvFile, latestFileName):
    """
    wordFile：要添加表格的word
    csvFile：要添加的表格
    latestFileName：最终生成的word文件的名称
    """
    doc = docx.Document(wordFile)
    index = 0
    """
    重构一下代码，这里性能太差
    """
    csvFile = multi_table_in_one_chart(csvFile)
    for csvIndex in csvFile:
        if not csvIndex:
            continue
        if len(csvIndex) <= 21:
            print("表格路径过短，已跳过，csv中标注路径为：" + csvIndex)
            continue
        if len(csvIndex) > 40:
            print("表格路径过长，已跳过，csv中标注路径为：" + csvIndex)
            continue
        csvIndex = csvIndex[3:]
        csvIndex = CURRENT_DIR + csvIndex
        csvIndex = csvIndex.replace(" ", "")

        with open(csvIndex, encoding="utf-8") as f:
            f_csv = csv.DictReader(f)
            ROWS = 0  # csv表格文件的行数
            csvKeyInotList = []
            csvValueInotList = []
            CSVColumn = 0  # csv表格的列，通过表格项总数/行数来计算得到，for循环之后更新
            for row in f_csv:
                # csv行数，此系数极为重要
                ROWS = ROWS + 1
                for key, value in row.items():
                    csvKeyInotList.append(key)
                    csvValueInotList.append(value)
                    CSVColumn = CSVColumn + 1
            CSVColumn = int(CSVColumn / ROWS)

        del csvKeyInotList[CSVColumn:]

        csvKeyInotList = np.array(csvKeyInotList)
        csvValueInotList = np.array(csvValueInotList).reshape(ROWS, CSVColumn)

        table = doc.add_table(rows=ROWS + 1, cols=CSVColumn, style='Table Grid')
        table.autofit = True
        table.alignment = WD_TABLE_ALIGNMENT.CENTER  # 表格居中对齐

        # 下面的代码是根据表格的列数控制每个表格的列的宽度
        if CSVColumn == 3:
            for i in range(0, CSVColumn - 1):
                for cell in table.columns[i].cells:
                    cell.width = Inches(1.4)
        elif CSVColumn == 6:
            for i in range(0, CSVColumn - 1):
                for cell in table.columns[i].cells:
                    cell.width = Inches(1.0)
                for cell in table.columns[5].cells:
                    cell.width = Inches(2.0)
        elif CSVColumn == 7:
            for i in range(0, CSVColumn - 1):
                for cell in table.columns[0].cells:
                    cell.width = Inches(2.4)

        for table_col in range(0, CSVColumn):
            table.rows[0].cells[table_col].text = csvKeyInotList[table_col]
            table.rows[0].cells[table_col].paragraphs[0].runs[0].font.size = Pt(12)
            for table_row in range(1, ROWS + 1):
                table.rows[table_row].cells[table_col].text = csvValueInotList[table_row - 1][table_col]
                table.rows[table_row].cells[table_col].paragraphs[0].runs[0].font.size = Pt(12)
                # set_cell_margins(table.rows[table_row].cells[table_col], right=0)  # 控制左右上下边距，但是唯独右好像不管用，故先注释

                # 首行灰色
        rows = table.rows[0]
        for cell in rows.cells:
            shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
            cell._tc.get_or_add_tcPr().append(shading_elm)

        # print("段落数:" + str(len(doc.paragraphs)))
        # print(doc.paragraphs[0])
        # 此处doc.paragraphs[]为经验值，试出来的

        # move_table_after函数极为重要，提供将新建的table移动到哪里的功能
        move_table_after(table, doc.paragraphs[7 + index])

        table_headers = doc.add_table(1, cols=1)
        cell = table_headers.cell(0, 0)
        # cell.line_spacing = 1.5

        table_name = ""
        if csvIndex[19] == "1":
            table_name = "各壳层星链卫星在轨总体情况"
        elif csvIndex[19] == "2":
            table_name = "各壳层星链卫星TLE数据情况"
        elif csvIndex[19] == "3":
            table_name = "星链卫星接近我国空间站情况"
        else:
            table_name = "卫星轨道升高幅度较大的统计"

        cell.paragraphs[0].add_run("附表 " + str(index + 1) + table_name).bold = True

        move_table_after(table_headers, doc.paragraphs[7 + index])

        index = index + 1
        delete_blank_pages(latestFileName)
        doc.save(latestFileName)


"""
# 该函数提供从csv中解析出table路径的功能，但是由于功能重复，且IO重复占用，暂时取消使用
def getTargetFile(file):
    # 获取要添加的csv文件
    csvList = []
    with open(file) as f:
        csvList = []
        f_csv = csv.DictReader(f)
        for key in f_csv:
            csvList.append(key["SourceUri"])
        return csvList
"""


def ProcessingThread(csv_list, counter):
    for index in csv_list:
        counter = counter + 1
        TABLE_PATH = FormatCSV(index)
        tpl = DocxTemplate(TEMPFILE)
        tpl.render(DIC)  # 渲染替换
        tpl.save(NEW_FILE_PATH + r"{}.docx".format(time[0]))

        add_table(NEW_FILE_PATH + r"{}.docx".format(time[0]), TABLE_PATH,
                  NEW_FILE_PATH + r"{}.docx".format(time[0]))
        delete_blank_pages(NEW_FILE_PATH + r"{}.docx".format(time[0]))
        print("Newly generated file： ", NEW_FILE_PATH + r"{}.docx".format(time[0]))
        print("Total number of documents: " + str(len(csv_list)) + "\t Generated: " + str(counter))


def predictTime(counter, tic, toc):
    seconds = ((toc - tic) / counter) * 1000
    m, s = divmod(seconds, 60)
    h, m = divmod(m, 60)
    print("if total number was 1000, it will cost:" + "%02d:%02d:%02d" % (h, m, s))


def get_parames_from_txt(params_file):
    with open(params_file, encoding='gbk') as f:
        data = f.read().split("\n")
        data_path = data[2].split("&&")[1] + '/rptdaily/'
        data_path = data_path.replace('\\', '/')
        print("data path: " + data_path)

    global CURRENT_DIR
    global TEMPFILE
    global CSV_PATH
    global NEW_FILE_PATH

    CURRENT_DIR = data_path
    TEMPFILE = os.getcwd() + '/新模版.docx'
    CSV_PATH = data_path + 'rptD_ONE/'
    NEW_FILE_PATH = data_path + 'output_doc/'


def main():
    get_parames_from_txt("PARAMS.txt")
    # return
    tic = t.perf_counter()
    global counter
    counter = 0
    EmptyDir(NEW_FILE_PATH)
    global csv_list
    csv_list = list_dir(CSV_PATH)
    ProcessingThread(csv_list, counter)
    print("----------------------------finish generating words!! ")

    # 这里不建议一起处理，会造成word软件进程冲突
    # word2pdf.doJob(NEW_FILE_PATH)

    toc = t.perf_counter()
    print(
        "----------------------------finish generating pdfs!! \n" + "well done " + f"time cost: {toc - tic:0.4f} seconds")


if __name__ == "__main__":
    main()
