# Import docx NOT python-docx
import docx

# Create an instance of a word document
doc = docx.Document()

# Add a Title to the document
doc.add_heading('GeeksForGeeks', 0)

# Table data in a form of list
data = (
    (1, 'Geek 1'),
    (2, 'Geek 2'),
    (3, 'Geek 3')
)

# Creating a table object
table = doc.add_table(rows=1, cols=2)

# Adding heading in the 1st row of the table
row = table.rows[0].cells
row[0].text = 'Id'
row[1].text = 'Name'

# Adding data from the list to the table
for id, name in data:
    # Adding a row and then adding data in it.
    row = table.add_row().cells
    row[0].text = str(id)
    row[1].text = name

# Adding style to a table
table.style = 'Colorful List'

# Now save the document to a location
doc.save('gfg.docx')
