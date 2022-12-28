# -*- coding : utf-8-*-
from docxtpl import DocxTemplate
import docx
import pandas as pd
import os
import csv
import numpy as np
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

CURRENT_DIR = os.getcwd() + '/'  # 获取当前的路径
TEMPFILE = "情报要报.docx"
CSV_PATH = os.getcwd() + "/data/"
NEW_FILE_PATH = os.getcwd() + "/new/"


def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def mkdir(path):
    path = path.strip()
    isExists = os.path.exists(path)
    if not isExists:
        os.makedirs(path.encode('utf-8'))
        return path
    else:
        return path


def list_dir(file_dir):
    list_csv = []
    dir_list = os.listdir(file_dir)
    for cur_file in dir_list:
        path = os.path.join(file_dir, cur_file)
        if os.path.isfile(path):
            dir_files = os.path.join(file_dir, cur_file)
        if os.path.splitext(path)[1] == '.csv':
            csv_file = os.path.join(file_dir, cur_file)
            list_csv.append(csv_file)
        if os.path.isdir(path):
            list_dir(path)
    return list_csv


def list_word_dir(file_dir):
    list_csv = []
    dir_list = os.listdir(file_dir)
    for cur_file in dir_list:
        path = os.path.join(file_dir, cur_file)
        if os.path.isfile(path):
            dir_files = os.path.join(file_dir, cur_file)
        if os.path.splitext(path)[1] == '.docx':
            csv_file = os.path.join(file_dir, cur_file)
            list_csv.append(csv_file)
        if os.path.isdir(path):
            list_dir(path)
    return list_csv


def makeDic(key, value):
    dic = {key: value}
    return dic


def EmptyDir(filepath):
    # 清空目标文件夹
    del_list = os.listdir(filepath)
    for f in del_list:
        file_path = os.path.join(filepath, f)
        if os.path.isfile(file_path):
            os.remove(file_path)


def ReFormatColumn(data, number, textContent, row, key):
    """
    data：字典中的值，将很多条目合并到一起的目标，入值最好为空字符串
    number：csv文件中内容条数索引
    content：csv文件中的内容，如biref和Content
    row:csv文件行数，以供遍历
    key:csv文件的键
    """
    for i in range(row):
        data += str(number[i])
        data += '. '
        data += ' '
        data += textContent[i].replace("xAyBzC01", "\n    ")
        data += '\n    '
    return makeDic(key, data)


def FormatCSV(csvfile):
    global DIC
    DIC = dict()

    global time
    data = pd.read_csv(csvfile)
    # encoding = 'gbk'
    time = data["Time"]
    rerange = data["Issue"]
    number = data["Item"]
    Type = data["Type"]
    Importance = data["Importance"]
    Brief = data["Brief"]
    Content = data["Content"]
    Publisher = data["Publisher"]
    SourceUri = data["SourceUri"]

    num = data.shape[0]  # 获取数据行数

    global Brief_dic
    Brief_data = str()
    Brief_dic = ReFormatColumn(Brief_data, number, Brief, num, "Brief")
    DIC.update(Brief_dic)

    global Content_dic
    Content_data = str()
    Content_dic = ReFormatColumn(Content_data, number, Content, num, "Content")
    DIC.update(Content_dic)

    Time_dic = {"Time": time[0]}
    DIC.update(Time_dic)

    Rerange_dic = {"期号": rerange[0]}
    DIC.update(Rerange_dic)

    End_dic = {"此页无正文": "（此页无正文）"}
    DIC.update(End_dic)
    # print("Final dictionary: ", '\n', DIC)


def delete_paragraph(paragraph):
    # 清除word中多余空行
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def DeleteBlankPages(xxx):
    # 清除word中空白页面
    myDoc = Document(xxx)
    for num, paragraphs in enumerate(myDoc.paragraphs):
        if paragraphs.text == "":
            # print('第{}段是空行段'.format(num))
            paragraphs.clear()  # 清除文字，并不删除段落，run也可以,
            delete_paragraph(paragraphs)


def PrintWord():
    file = docx.Document(NEW_FILE_PATH)
    print("段落数:" + str(len(file.paragraphs)))  # 段落数为13，每个回车隔离一段

    # 输出段落编号及段落内容
    for i in range(len(file.paragraphs)):
        if len(file.paragraphs[i].text.replace(' ', '')) > 4:
            print("第" + str(i) + "段的内容是：" + file.paragraphs[i].text)


def translate(str, old, new):
    str.replace(old, new)


def AddTable(wordFile, csvFile, latestFileName):
    """
    wordFile：要添加表格的word
    csvFile：要添加的表格
    latestFileName：最终生成的word文件的名称
    """

    index = 0
    for csvFile in csvFile:
        if not csvFile:
            continue
        if len(csvFile) == 0:
            continue
        if csvFile[1] == "t":
            continue

        doc = docx.Document(wordFile)
        csvFile = csvFile[1:]

        with open(csvFile) as f:
            f_csv = csv.DictReader(f)
            ROWS = 0
            csvKeyInotList = []
            csvValueInotList = []
            for row in f_csv:
                # csv行数，此系数极为重要
                ROWS = ROWS + 1
                for key, value in row.items():
                    csvKeyInotList.append(key)
                    csvValueInotList.append(value)
        del csvKeyInotList[11:]
        csvKeyInotList = np.array(csvKeyInotList)
        csvValueInotList = np.array(csvValueInotList).reshape(ROWS, 11)

        table = doc.add_table(rows=ROWS + 1, cols=11, style='Table Grid')
        table.autofit = True
        table_header = table.rows[0].cells

        for col_x in range(0, 11):
            table_header[col_x].text = csvKeyInotList[col_x]
        for row in range(1, ROWS + 1):
            for col_y in range(0, 11):
                table_x = table.rows[row].cells
                table_x[col_y].text = csvValueInotList[row - 1][col_y]

        # 首行灰色
        rows = table.rows[0]
        for cell in rows.cells:
            shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
            cell._tc.get_or_add_tcPr().append(shading_elm)

        # print("段落数:" + str(len(doc.paragraphs)))
        # print(doc.paragraphs[0])
        # 此处doc.paragraphs[]为经验值，试出来的

        # move_table_after函数极为重要，提供将新建的table移动到哪里的功能
        move_table_after(table, doc.paragraphs[11 + index])

        table_headers = doc.add_table(1, cols=1)
        cell = table_headers.cell(0, 0)
        cell.line_spacing = 0.5

        table_name = ""
        if csvFile[29] == "1":
            table_name = "卫星轨道降低幅度较大的统计"
        elif csvFile[29] == "2":
            table_name = "卫星轨道降低幅度较小的统计"
        elif csvFile[29] == "3":
            table_name = "卫星轨道升高幅度较小的统计"
        else:
            "卫星轨道升高幅度较大的统计"

        cell.paragraphs[0].add_run("附表 " + str(index + 1) + table_name).bold = True
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 水平居中

        move_table_after(table_headers, doc.paragraphs[11 + index])

        index = index + 1
        DeleteBlankPages(latestFileName)

        doc.save(latestFileName)


def getTargetFile(file):
    # 获取要添加的csv文件
    csvList = []
    with open(file) as f:
        csvList = []
        f_csv = csv.DictReader(f)
        for key in f_csv:
            csvList.append(key["SourceUri"])
        return csvList


def main():
    EmptyDir(NEW_FILE_PATH)
    csv_list = list_dir(CSV_PATH)
    print("csv_fils: ", '\n', csv_list)

    for index in csv_list:
        FormatCSV(index)
        tpl = DocxTemplate(CURRENT_DIR + TEMPFILE)
        tpl.render(DIC)  # 渲染替换
        tpl.save(NEW_FILE_PATH + r"{}.docx".format(time[0]))

        # 添加table，第一个参数为word文档，第二个参数为对应解析出来的csv文件
        print(index)
        AddTable(NEW_FILE_PATH + r"{}.docx".format(time[0]), getTargetFile(index),
                 NEW_FILE_PATH + r"{}.docx".format(time[0]))

        print("Newly generated files： ", NEW_FILE_PATH + r"{}.docx".format(time[0]))
        DeleteBlankPages(NEW_FILE_PATH + r"{}.docx".format(time[0]))


if __name__ == "__main__":
    main()
