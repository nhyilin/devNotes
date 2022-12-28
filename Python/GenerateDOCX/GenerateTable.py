# Import docx NOT python-docx
import docx

def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)

dict = {'name': 'Zara', 'age': 7, 'class': 'First'}

def main():
    # /Users/peiyilin/dev/devNotes/Python/GenerateDOCX/test/表格生成测试.docx
    doc = docx.Document("/Users/peiyilin/dev/devNotes/Python/GenerateDOCX/test/表格生成测试.docx")

    # Create an instance of a word document
    # doc = docx.Document()
    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph('')
    # doc.add_heading('表标题', 0)
    #
    # # Table data in a form of list
    data = (
        (1, 'Geek 1'),
        (2, 'Geek 2'),
        (3, 'Geek 3')
    )
    #
    # # Creating a table object
    table = doc.add_table(rows=1, cols=2, style='Table Grid')
    #
    # Adding heading in the 1st row of the table
    row = table.rows[0].cells
    row[0].text = 'Id'
    row[1].text = 'Name'

    # Adding data from the list to the table
    for id, name in data:
        # Adding a row and then adding data in it.
        row = table.add_row().cells
        # Converting id to string as table can only take string input
        row[0].text = str(id)
        row[1].text = name

    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph('')
    # Now save the document to a location

    print("段落数:" + str(len(doc.paragraphs)))
    print(doc.paragraphs[0])
    move_table_after(table, doc.paragraphs[20])

    doc.save('gfg.docx')


if __name__ == "__main__":
    main()
