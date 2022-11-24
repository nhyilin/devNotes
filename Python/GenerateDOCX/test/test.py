from docx import Document
import time

ocaltime = time.localtime(time.time())
time = time.strftime("%H_%M_%S", time.localtime())


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


myDocument = Document('./2022.11.19.docx')
for num, paragraphs in enumerate(myDocument.paragraphs):
    if paragraphs.text == "":
        print('第{}段是空行段'.format(num))
        paragraphs.clear()  # 清除文字，并不删除段落，run也可以,
        delete_paragraph(paragraphs)

# p1 = myDocument.add_paragraph('')
# # for x in range(7):
#     # h2 = myDocument.add_heading('\n', level=2)
# p1 = myDocument.add_paragraph('uper')
# myDocument.add_page_break()
# p1 = myDocument.add_paragraph('downer')
# p1 = myDocument.add_paragraph('')
# for x in range(7):
#     h2 = myDocument.add_heading('\n', level=2)

# table = myDocument.add_table(rows=3, cols=2)
# row1 = table.rows[0]
# row1.cells[0].text = '报： '
# row2 = table.rows[1]
# row2.cells[0].text = '送： '
# row3 = table.rows[2]
# row3.cells[0].text = '责任编辑：'
# row3.cells[1].text = '电话：'
# for x in range(6):


myDocument.save("xxx" + time + '.docx')
